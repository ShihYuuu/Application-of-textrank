{
 "cells": [
  {
   "cell_type": "code",
   "execution_count": 3,
   "metadata": {},
   "outputs": [],
   "source": [
    "import pandas as pd"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 4,
   "metadata": {},
   "outputs": [],
   "source": [
    "df = pd.read_excel(\"company_100.xlsx\", header=0)"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 5,
   "metadata": {},
   "outputs": [],
   "source": [
    "company_names = df['Name'].values.tolist()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 6,
   "metadata": {},
   "outputs": [],
   "source": [
    "company_article = df['Profile'].dropna().values.tolist()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 7,
   "metadata": {},
   "outputs": [],
   "source": [
    "company_source = df['Source'].values.tolist()"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 8,
   "metadata": {},
   "outputs": [],
   "source": [
    "import numpy as np\n",
    "import pandas as pd\n",
    "import re\n",
    "import networkx as nx\n",
    "from sklearn.feature_extraction.text import CountVectorizer\n",
    "from sklearn.feature_extraction.text import TfidfVectorizer\n",
    "from sklearn.metrics.pairwise import cosine_similarity"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 39,
   "metadata": {},
   "outputs": [],
   "source": [
    "from ckiptagger import data_utils, construct_dictionary, WS, POS, NER\n",
    "ws = WS(\"./data\") #WS段詞 POS詞性標註 NER實體辨識"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 40,
   "metadata": {},
   "outputs": [],
   "source": [
    "# 讀入停用詞檔\n",
    "stopWords=[]\n",
    "\n",
    "with open('stopWords.txt', 'r', encoding='UTF-8') as file:\n",
    "    for data in file.readlines():\n",
    "        data = data.strip()\n",
    "        stopWords.append(data)"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 41,
   "metadata": {},
   "outputs": [],
   "source": [
    "def read_article(file_name):\n",
    "    file = open(file_name, \"r\", encoding=\"utf-8\")\n",
    "    full_text = \"\"\n",
    "    for line in file.readlines():\n",
    "        line = line.strip('\\n')\n",
    "        full_text+=line  \n",
    "        \n",
    "    article = re.split(\"[。，]\", full_text) \n",
    "    article.pop(-1)\n",
    "    \n",
    "    return article"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 42,
   "metadata": {},
   "outputs": [],
   "source": [
    "def read_article_from_list(article):\n",
    "    full_text = article\n",
    "#     full_text = full_text.replace(\"\\n\",\"\")\n",
    "    full_text = full_text.replace(\"經營理念\",\"\")\n",
    "    article = re.split(\"[。|\\n]\", full_text)\n",
    "    for i in article:\n",
    "        i.replace(\"\\n\",\"\")\n",
    "    article.pop(-1)\n",
    "    \n",
    "    return article"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 43,
   "metadata": {},
   "outputs": [],
   "source": [
    "def build_similarity_matrix(article, stopWords):\n",
    "    word_sentence_list = ws(\n",
    "        article,\n",
    "        # sentence_segmentation=True, # To consider delimiters\n",
    "        # segment_delimiter_set = {\",\", \"。\", \":\", \"?\", \"!\", \";\"}),\n",
    "        # This is the defualt set of delimiters\n",
    "        # recommend_dictionary = dictionary1, # words in this dictionary are encouraged\n",
    "        # coerce_dictionary = dictionary2, # words in this dictionary are forced\n",
    "    )\n",
    "    \n",
    "    corpus = []\n",
    "    for segments in word_sentence_list:\n",
    "        remainderWords = list(filter(lambda a: a not in stopWords and a != '\\n', segments))\n",
    "        corpus.append(' '.join(remainderWords))\n",
    "    \n",
    "    vectorizer = CountVectorizer()\n",
    "    X = vectorizer.fit_transform(corpus)\n",
    "#     print(\"Bag-of-words:\\n\", vectorizer.get_feature_names())\n",
    "\n",
    "#     tfidf = TfidfVectorizer()\n",
    "#     X = tfidf.fit_transform(corpus)\n",
    "    \n",
    "    cs = cosine_similarity(X)\n",
    "    # make same doc similarity score=0\n",
    "    for i in range(len(cs)):\n",
    "        for j in range(len(cs)):\n",
    "            if i==j:\n",
    "                cs[i][j]=0\n",
    "                \n",
    "#     print(\"\\nCosine similarity matrix:\\n\", cs)\n",
    "    return cs"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 44,
   "metadata": {},
   "outputs": [],
   "source": [
    "def create_article_summary(index, text):\n",
    "    print(\"-\"*50,index,\"-\"*50)\n",
    "    \n",
    "    # Step 1 - Read txt file and cut sentences with。\n",
    "    article = read_article_from_list(text)\n",
    "    article = [k for k in article if k.strip()!='']\n",
    "    # print(\"# of sentences in article: \", len(article))\n",
    "\n",
    "\n",
    "    # Step 2 - Generate Similary Martix across sentences\n",
    "    sentence_similarity_martix = build_similarity_matrix(article, stopWords)\n",
    "\n",
    "\n",
    "    # Step 3 - Rank sentences in similarity martix\n",
    "    sentence_similarity_graph = nx.from_numpy_array(sentence_similarity_martix)\n",
    "    scores = nx.pagerank(sentence_similarity_graph)\n",
    "\n",
    "\n",
    "    # Step 4 - Sort the rank and pick top sentences\n",
    "    ranked_sentence = sorted(((scores[i],s) for i,s in enumerate(article)), reverse=True)    \n",
    "    # print(\"Indexes of top ranked_sentence order are \\n\", ranked_sentence)\n",
    "    \n",
    "    import math\n",
    "    # Step 5 - Decide top_n\n",
    "    top_n = math.ceil(len(article)/2)\n",
    "    print(\"article length: %d\\ntop_n: %d\"%(len(article), top_n))\n",
    "    \n",
    "    summarize_text = []\n",
    "    for i in range(top_n):\n",
    "        summarize_text.append(\"\".join(ranked_sentence[i][1]))\n",
    "        \n",
    "    # Step 5 - Offcourse, output the summarize texr\n",
    "    # print(\"Summarize Text:\\n\", \"\\n\".join(summarize_text))\n",
    "\n",
    "\n",
    "    ordered_summarize_text = \"\"\n",
    "    for sent in article:\n",
    "        if sent in summarize_text:\n",
    "    #         print(sent+'。')\n",
    "            ordered_summarize_text+=(sent+'。\\n')\n",
    "    \n",
    "    print(ordered_summarize_text)\n",
    "    \n",
    "    return ordered_summarize_text"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 9,
   "metadata": {
    "jupyter": {
     "outputs_hidden": true
    }
   },
   "outputs": [
    {
     "ename": "AttributeError",
     "evalue": "'str' object has no attribute 'dtypes'",
     "output_type": "error",
     "traceback": [
      "\u001b[1;31m---------------------------------------------------------------------------\u001b[0m",
      "\u001b[1;31mAttributeError\u001b[0m                            Traceback (most recent call last)",
      "\u001b[1;32m<ipython-input-9-37278870f5a8>\u001b[0m in \u001b[0;36m<module>\u001b[1;34m\u001b[0m\n\u001b[0;32m      2\u001b[0m \u001b[1;33m\u001b[0m\u001b[0m\n\u001b[0;32m      3\u001b[0m \u001b[1;32mfor\u001b[0m \u001b[0mindex\u001b[0m\u001b[1;33m,\u001b[0m \u001b[0mtext\u001b[0m \u001b[1;32min\u001b[0m \u001b[0menumerate\u001b[0m\u001b[1;33m(\u001b[0m\u001b[0mcompany_article\u001b[0m\u001b[1;33m)\u001b[0m\u001b[1;33m:\u001b[0m\u001b[1;33m\u001b[0m\u001b[1;33m\u001b[0m\u001b[0m\n\u001b[1;32m----> 4\u001b[1;33m     \u001b[0mtext\u001b[0m\u001b[1;33m.\u001b[0m\u001b[0mdtypes\u001b[0m\u001b[1;33m\u001b[0m\u001b[1;33m\u001b[0m\u001b[0m\n\u001b[0m\u001b[0;32m      5\u001b[0m     \"\"\"\n\u001b[0;32m      6\u001b[0m     \u001b[1;32mtry\u001b[0m\u001b[1;33m:\u001b[0m\u001b[1;33m\u001b[0m\u001b[1;33m\u001b[0m\u001b[0m\n",
      "\u001b[1;31mAttributeError\u001b[0m: 'str' object has no attribute 'dtypes'"
     ]
    }
   ],
   "source": [
    "company_summary = []\n",
    "\n",
    "for index, text in enumerate(company_article):\n",
    "    \n",
    "    try:\n",
    "        summary = create_article_summary(index, text)\n",
    "    except:\n",
    "        summary = text\n",
    "    company_summary.append([company_names[index], text, summary])\n",
    "    "
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 70,
   "metadata": {},
   "outputs": [],
   "source": [
    "def output_docx(company_summary, k):\n",
    "    from docx import Document\n",
    "    document = Document()\n",
    "    document.add_heading('Company Profile', level=1)\n",
    "    paragraph = document.add_paragraph(company_summary[1])\n",
    "    document.add_page_break()\n",
    "    document.add_heading('Company Profile Summary', level=1)\n",
    "    paragraph = document.add_paragraph(company_summary[2])\n",
    "    \n",
    "    if(type(company_source[k]) != float):\n",
    "        document.add_heading('Company Profile Source', level=1)\n",
    "        paragraph = document.add_paragraph(company_source[k])\n",
    "    document.save('./summaries/%s.docx' %(company_summary[0]))"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 73,
   "metadata": {
    "scrolled": true
   },
   "outputs": [],
   "source": [
    "k = 0\n",
    "for i in company_summary:\n",
    "    try:\n",
    "    output_docx(i, k)\n",
    "    k += 1\n",
    "    except:\n",
    "        print(i[0])"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 72,
   "metadata": {},
   "outputs": [],
   "source": [
    "def output_xlsx(company_summary):\n",
    "    from pandas.core.frame import DataFrame\n",
    "    df_output = DataFrame(company_summary)\n",
    "    df_output.columns = ['name', 'profile', 'summary']\n",
    "    df_output.to_excel(\"company_summary.xlsx\", encoding=\"utf8\", index=None)"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": 2,
   "metadata": {},
   "outputs": [
    {
     "name": "stdout",
     "output_type": "stream",
     "text": [
      "This application is used to convert notebook files (*.ipynb) to various other\n",
      "formats.\n",
      "\n",
      "WARNING: THE COMMANDLINE INTERFACE MAY CHANGE IN FUTURE RELEASES.\n",
      "\n",
      "Options\n",
      "\n",
      "-------\n",
      "\n",
      "\n",
      "\n",
      "Arguments that take values are actually convenience aliases to full\n",
      "Configurables, whose aliases are listed on the help line. For more information\n",
      "on full configurables, see '--help-all'.\n",
      "\n",
      "\n",
      "--debug\n",
      "\n",
      "    set log level to logging.DEBUG (maximize logging output)\n",
      "\n",
      "--generate-config\n",
      "\n",
      "    generate default config file\n",
      "\n",
      "-y\n",
      "\n",
      "    Answer yes to any questions instead of prompting.\n",
      "\n",
      "--execute\n",
      "\n",
      "    Execute the notebook prior to export.\n",
      "\n",
      "--allow-errors\n",
      "\n",
      "    Continue notebook execution even if one of the cells throws an error and include the error message in the cell output (the default behaviour is to abort conversion). This flag is only relevant if '--execute' was specified, too.\n",
      "\n",
      "--stdin\n",
      "\n",
      "    read a single notebook file from stdin. Write the resulting notebook with default basename 'notebook.*'\n",
      "\n",
      "--stdout\n",
      "\n",
      "    Write notebook output to stdout instead of files.\n",
      "\n",
      "--inplace\n",
      "\n",
      "    Run nbconvert in place, overwriting the existing notebook (only \n",
      "    relevant when converting to notebook format)\n",
      "\n",
      "--clear-output\n",
      "\n",
      "    Clear output of current file and save in place, \n",
      "    overwriting the existing notebook.\n",
      "\n",
      "--no-prompt\n",
      "\n",
      "    Exclude input and output prompts from converted document.\n",
      "\n",
      "--no-input\n",
      "\n",
      "    Exclude input cells and output prompts from converted document. \n",
      "    This mode is ideal for generating code-free reports.\n",
      "--log-level=<Enum> (Application.log_level)\n",
      "\n",
      "    Default: 30\n",
      "\n",
      "    Choices: (0, 10, 20, 30, 40, 50, 'DEBUG', 'INFO', 'WARN', 'ERROR', 'CRITICAL')\n",
      "\n",
      "    Set the log level by value or name.\n",
      "\n",
      "--config=<Unicode> (JupyterApp.config_file)\n",
      "\n",
      "    Default: ''\n",
      "\n",
      "    Full path of a config file.\n",
      "\n",
      "--to=<Unicode> (NbConvertApp.export_format)\n",
      "\n",
      "    Default: 'html'\n",
      "\n",
      "    The export format to be used, either one of the built-in formats\n",
      "\n",
      "    ['asciidoc', 'custom', 'html', 'latex', 'markdown', 'notebook', 'pdf',\n",
      "\n",
      "    'python', 'rst', 'script', 'slides'] or a dotted object name that represents\n",
      "\n",
      "    the import path for an `Exporter` class\n",
      "\n",
      "--template=<Unicode> (TemplateExporter.template_file)\n",
      "\n",
      "    Default: ''\n",
      "\n",
      "    Name of the template file to use\n",
      "\n",
      "--writer=<DottedObjectName> (NbConvertApp.writer_class)\n",
      "\n",
      "    Default: 'FilesWriter'\n",
      "\n",
      "    Writer class used to write the  results of the conversion\n",
      "\n",
      "--post=<DottedOrNone> (NbConvertApp.postprocessor_class)\n",
      "\n",
      "    Default: ''\n",
      "\n",
      "    PostProcessor class used to write the results of the conversion\n",
      "\n",
      "--output=<Unicode> (NbConvertApp.output_base)\n",
      "\n",
      "    Default: ''\n",
      "\n",
      "    overwrite base name use for output files. can only be used when converting\n",
      "\n",
      "    one notebook at a time.\n",
      "\n",
      "--output-dir=<Unicode> (FilesWriter.build_directory)\n",
      "\n",
      "    Default: ''\n",
      "\n",
      "    Directory to write output(s) to. Defaults to output to the directory of each\n",
      "\n",
      "    notebook. To recover previous default behaviour (outputting to the current\n",
      "\n",
      "    working directory) use . as the flag value.\n",
      "\n",
      "--reveal-prefix=<Unicode> (SlidesExporter.reveal_url_prefix)\n",
      "\n",
      "    Default: ''\n",
      "\n",
      "    The URL prefix for reveal.js (version 3.x). This defaults to the reveal CDN,\n",
      "\n",
      "    but can be any url pointing to a copy  of reveal.js.\n",
      "\n",
      "    For speaker notes to work, this must be a relative path to a local  copy of\n",
      "\n",
      "    reveal.js: e.g., \"reveal.js\".\n",
      "\n",
      "    If a relative path is given, it must be a subdirectory of the current\n",
      "\n",
      "    directory (from which the server is run).\n",
      "\n",
      "    See the usage documentation\n",
      "\n",
      "    (https://nbconvert.readthedocs.io/en/latest/usage.html#reveal-js-html-\n",
      "\n",
      "    slideshow) for more details.\n",
      "\n",
      "--nbformat=<Enum> (NotebookExporter.nbformat_version)\n",
      "\n",
      "    Default: 4\n",
      "\n",
      "    Choices: [1, 2, 3, 4]\n",
      "\n",
      "    The nbformat version to write. Use this to downgrade notebooks.\n",
      "\n",
      "To see all available configurables, use `--help-all`\n",
      "\n",
      "Examples\n",
      "--------\n",
      "\n",
      "    The simplest way to use nbconvert is\n",
      "    \n",
      "    > jupyter nbconvert mynotebook.ipynb\n",
      "    \n",
      "    which will convert mynotebook.ipynb to the default format (probably HTML).\n",
      "    \n",
      "    You can specify the export format with `--to`.\n",
      "    Options include ['asciidoc', 'custom', 'html', 'latex', 'markdown', 'notebook', 'pdf', 'python', 'rst', 'script', 'slides'].\n",
      "    \n",
      "    > jupyter nbconvert --to latex mynotebook.ipynb\n",
      "    \n",
      "    Both HTML and LaTeX support multiple output templates. LaTeX includes\n",
      "    'base', 'article' and 'report'.  HTML includes 'basic' and 'full'. You\n",
      "    can specify the flavor of the format used.\n",
      "    \n",
      "    > jupyter nbconvert --to html --template basic mynotebook.ipynb\n",
      "    \n",
      "    You can also pipe the output to stdout, rather than a file\n",
      "    \n",
      "    > jupyter nbconvert mynotebook.ipynb --stdout\n",
      "    \n",
      "    PDF is generated via latex\n",
      "    \n",
      "    > jupyter nbconvert mynotebook.ipynb --to pdf\n",
      "    \n",
      "    You can get (and serve) a Reveal.js-powered slideshow\n",
      "    \n",
      "    > jupyter nbconvert myslides.ipynb --to slides --post serve\n",
      "    \n",
      "    Multiple notebooks can be given at the command line in a couple of \n",
      "    different ways:\n",
      "    \n",
      "    > jupyter nbconvert notebook*.ipynb\n",
      "    > jupyter nbconvert notebook1.ipynb notebook2.ipynb\n",
      "    \n",
      "    or you can specify the notebooks list in a config file, containing::\n",
      "    \n",
      "        c.NbConvertApp.notebooks = [\"my_notebook.ipynb\"]\n",
      "    \n",
      "    > jupyter nbconvert --config mycfg.py\n",
      "\n"
     ]
    },
    {
     "name": "stderr",
     "output_type": "stream",
     "text": [
      "[NbConvertApp] WARNING | pattern 'to' matched no files\n",
      "[NbConvertApp] WARNING | pattern 'script' matched no files\n",
      "[NbConvertApp] WARNING | pattern 'Text_Rank_summary.ipynb' matched no files\n"
     ]
    }
   ],
   "source": [
    "!jupyter nbconvert -- to script Text_Rank_summary.ipynb"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": []
  }
 ],
 "metadata": {
  "kernelspec": {
   "display_name": "Python 3",
   "language": "python",
   "name": "python3"
  },
  "language_info": {
   "codemirror_mode": {
    "name": "ipython",
    "version": 3
   },
   "file_extension": ".py",
   "mimetype": "text/x-python",
   "name": "python",
   "nbconvert_exporter": "python",
   "pygments_lexer": "ipython3",
   "version": "3.7.4"
  }
 },
 "nbformat": 4,
 "nbformat_minor": 4
}
